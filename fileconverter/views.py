from .forms import UploadFileForm
from .forms import ChangeRequestFileForm
import openpyxl
from docx.shared import Cm
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.views.decorators.csrf import csrf_exempt
from docx.shared import Inches
import os
import logging
from docx.image.exceptions import UnexpectedEndOfFileError

logger = logging.getLogger(__name__)

# Хувийн хэрэгт шаардлагатай баганын тодорхойлолт
# index: (нэр, заавал эсэх)
HUVIIN_HEREG_COLUMNS = {
    1:  ("Дахин давтагдашгүй дугаар", True),
    2:  ("Нэрийн зургийн индекс",     True),
    3:  ("Газар зүйн нэр",            True),
    4:  ("Төрөл",                     True),
    5:  ("Дэвсгэр нэр / Ангилал",     False),
    9:  ("Байр зүйн зургийн нэрлэвэр",False),
    10: ("1:100 000 зурагт",          False),
    13: ("Гарал үүсэл",               False),
    14: ("Өргөрөг 1",                 True),
    15: ("Уртраг 1",                  True),
    16: ("Өргөрөг 2",                 False),
    17: ("Уртраг 2",                  False),
    18: ("Аймаг/сум/баг",             True),
    19: ("Байрлал тайлбар",           False),
    20: ("Актын дугаар",              False),
}

def validate_row(row, row_num):
    """Мөрийн баганын утгуудыг шалгаж log хийнэ. True буцаана хэрэв заавал талбарууд бүгд байвал."""
    ok = True
    for idx, (name, required) in HUVIIN_HEREG_COLUMNS.items():
        val = row[idx] if idx < len(row) else None
        if val is None or str(val).strip() == "":
            if required:
                logger.warning("  [Мөр %d] ❌ ЗААВАЛ багана хоосон: col[%d] = '%s'", row_num, idx, name)
                ok = False
            else:
                logger.info("  [Мөр %d]    Хоосон (заавал биш): col[%d] = '%s'", row_num, idx, name)
        else:
            logger.info("  [Мөр %d] ✓ col[%d] '%s' = %s", row_num, idx, name, str(val)[:60])
    return ok

PAGE_WIDTH_INCHES = 8.5

EXPORT_PATHS  = [
"E:\\5 sum zurag\\bairshil\\Gurvantes\\42",
"E:\\5 sum zurag\\bairshil\\Gurvantes\\ulamjlalt",
"E:\\5 sum zurag\\bairshil\\Gurvantes\\bilge\\42",
"E:\\5 sum zurag\\bairshil\\Gurvantes\\bilge\\уламжлалт",
"E:\\5 sum zurag\\bairshil\\Gurvantes\\өөрчлөлт",
"E:\\5 sum zurag\\bairshil\\Gurvantes",
]


IMAGE_PATHS = [ 
    # 'E:\\5 sum zurag\\photos\\bayndalai',
    # 'E:\\5 sum zurag\\photos\\Bulgan',
    # 'E:\\5 sum zurag\\photos\\Noyn',
    # 'E:\\5 sum zurag\\photos\\Sevrei',
    'E:\\5 sum zurag\\photos\\Gurvantes',
]

def find_image_path(name):
    image_name = "%s.jpg" % name
    for i in IMAGE_PATHS:
        image_path = os.path.join(i, image_name)
        if os.path.exists(image_path):
            return image_path
        words = name.split()  # Split the string into a list of words
        words = words[:-1]  # Remove the last word by slicing
        words = ' '.join(words)
        if words == "":
            continue  
        image_name = "%s.jpg" % words
        image_path = os.path.join(i, image_name)
        if os.path.exists(image_path):
            return image_path
        words = words.split()  # Split the string into a list of words
        words = words[:-1]  # Remove the last word by slicing
        words = ' '.join(words)  # Join the remaining words back into a string
        if words == "":
            continue
        image_name = "%s.jpg" % words
        image_path = os.path.join(i, image_name)
        if os.path.exists(image_path):
            return image_path
    return False

def find_export_image_path(name, index):
    image_name = "%s %s.jpg" % (name, index)
    for i in EXPORT_PATHS:
        image_path = os.path.join(i, image_name)
        if os.path.exists(image_path):
            return image_path
    return False

def handle_uploaded_file(f):
    logger.info("=== handle_uploaded_file эхлэв ===")
    wb = openpyxl.load_workbook(f)
    ws = wb.active
    logger.info("Excel файл нээгдлээ. Sheet: %s, max_row=%s, max_col=%s",
                ws.title, ws.max_row, ws.max_column)
    doc = Document()
    static_values = [
    {
        "first": "Газар зүйн нэр /монгол, латин галиг/",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийн дахин давтагдашгүй дугаар",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийн гарал, үүсэл",
        "second": "Уламжлалт нэр"
    },
    {
        "first": "Газар зүйн нэрийн төрөл /дэвсгэр нэр/",
        "second": ""
    },
    {
        "first": "Харьяалагдах аймаг, сум, баг",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийн ерөнхий байрлал, тайлбар",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийн солбицол, UTM, 48-р бүс",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийн орших 1:25 000-ны масштабтай байр зүйн зургийн нэрлэвэр",
        "second": ""
    },
    {
        "first": "",
        "second": ""
    },
    {
        "first": "Бусад газрын зурагт зөв, ижил хэрэглэгдсэн байдал",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийг баталгаажуулсан актын нэр, дугаар, огноо",
        "second": "Сумын иргэдийн төлөөлөгчдийн ...-хурлаар дэмжигдсэн ."
    },
    {
        "first": "Өөрчлөлт орсон эсэх, шалтгаан",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийн байршлын зураг",
        "second": ""
    },
    {
        "first": "",
        "second": ""
    },
    {
        "first": "Газар зүйн нэрийн тодруулалтын үеийн нотлох баримт:",
        "second": "Аудио, видео бичлэг: □ \nТэмдэглэл:    □        Фото зураг:   □"
    },
    {
        "first": "Газар зүйн нэрийг тодруулсан иргэн, хуулийн этгээд",
        "second": "“Инженер геодези” ХХК-ны инженер:\nМУ-ын зөвлөх инженер Д.Оюунчимэг\nИнженер: Э.Ануун, Н.Бумчин"
    },
    {
        "first": "Газар зүйн нэрийг тодруулсан газарчин /орон нутгийн/",
        "second": "Н.Очирваань, багийн өндөр настан\nЭ.Эрдэнэтунгалаг, газрын даамал"
    },
    {
        "first": "Газар зүйн нэрийн хувийн хэрэг хөтөлсөн:",
        "second": "/2024 оны 05-р сарын 15-ны өдөр/"
    }]

    all_rows = list(ws.iter_rows(min_row=3, values_only=True))
    logger.info("Excel-ийн нийт мөр (header-гүй): %d", len(all_rows))

    # Өгөгдөл эхэлсний дараа анхны хоосон мөр дээр зогсоно
    # col[3] = Газар зүйн нэр — энэ хоосон бол өгөгдөл дууссан гэж үзнэ
    data_rows = []
    for i, row in enumerate(all_rows):
        name_val = row[3] if len(row) > 3 else None
        if name_val is None or str(name_val).strip() == "":
            if data_rows:
                logger.info("Excel мөр %d — хоосон мөр, боловсруулалт зогслоо.", i + 3)
                break
            else:
                logger.info("Excel мөр %d — өгөгдлийн өмнөх хоосон мөр, алгасав.", i + 3)
        else:
            data_rows.append((i + 3, row))  # (excel_row_num, row)

    logger.info("Боловсруулах мөр: %d", len(data_rows))

    if len(data_rows) == 0:
        raise ValueError("Боловсруулах өгөгдөл олдсонгүй. D баганад (Газар зүйн нэр) утга байхгүй байна.")

    # Validation: заавал талбар дутуу мөрүүдийг цуглуулах
    errors = []
    for row_num, row in data_rows:
        logger.info("--- Excel мөр %d шалгаж байна ---", row_num)
        for ci, val in enumerate(row):
            if val is not None and str(val).strip():
                logger.info("  col[%d] = %s", ci, str(val)[:80])
        is_valid = validate_row(row, row_num)
        if not is_valid:
            missing = [
                name for idx, (name, req) in HUVIIN_HEREG_COLUMNS.items()
                if req and (idx >= len(row) or row[idx] is None or str(row[idx]).strip() == "")
            ]
            errors.append("Excel мөр %d: %s" % (row_num, ", ".join(missing)))

    if errors:
        logger.warning("Validation алдаа: %d мөрт дутуу талбар байна", len(errors))
        raise ValueError("Дараах мөрүүдэд заавал талбар дутуу байна:\n" + "\n".join(errors))

    for row_num, row in data_rows:
        logger.info("--- Excel мөр %d боловсруулж байна ---", row_num)

        title = doc.add_paragraph('')
        run = title.add_run('Газар зүйн нэрийн хувийн хэрэг')
        run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=18, cols=10)
        table.style = 'Table Grid'
        table.autofit = False
        total_width_cm = 20  # Example total width for table
        table.columns[0].width = Cm(total_width_cm * 0.4)
        for i, static_value in enumerate(static_values):
            image_path = find_image_path(str(row[3]))
            export_path = find_export_image_path(str(row[3]), str(row[2]))
            row_cells = table.rows[i].cells
            if i != 12:
                row_cells[0].text = static_value.get("first", "")
                horizontal = table.cell(i, 0).merge(table.cell(i, 3))
                if i == 8:
                    horizontal.merge(table.cell(i+1, 3))
                elif i == 13:
                    paragraph = row_cells[0].paragraphs[0]
                    run = paragraph.add_run()
                    if image_path:
                        print("image_path found", image_path)
                        try:
                            run.add_picture(image_path, width=Inches(3.0), height=Inches(3.0))
                        except UnexpectedEndOfFileError:
                            print(f"Error: The image file at {image_path} is corrupted or malformed.")
                        except Exception as e:
                            print(f"An error occurred: {e}")
            if i == 0:
                row_cells[4].text = (str(row[3]) if row[3] is not None else '')
                row_cells[7].text = str(row[4]) if row[4] is not None else ''
                table.cell(i, 4).merge(table.cell(i, 6))
                table.cell(i, 7).merge(table.cell(i, 9))
            elif i == 1:
                row_cells[4].text = (str(row[1]) if row[1] is not None else '')
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 2:
                row_cells[4].text = (str(row[13]) if row[13] is not None else '')
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 3:
                row_cells[4].text = str(row[5]) if row[5] is not None else ''
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 4:
                row_cells[4].text = (str(row[18]) if row[18] is not None else '')
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 5:
                row_cells[4].text = (str(row[19]) if row[19] is not None else '')
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 6:
                longitude1 = "1.Өргөрөг: " + ("" if not row[14] else str(row[14]))
                latitude1 = "\nУртраг: " + ("" if not row[15] else str(row[15]))

                longitude2 = "\n2.Өргөрөг: " + ("" if not row[16] else str(row[16]))
                latitude2 = "\nУртраг: " + ("" if not row[17] else str(row[17]))
                coordinate = longitude1 + latitude1
                if row[16]:
                    coordinate = coordinate + longitude2 + latitude2
                row_cells[4].text = coordinate
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 7:
                row_cells[4].text = str(row[9]) if row[9] is not None else ''
                table.cell(i, 4).merge(table.cell(i, 5))
                row_cells[6].text = "Нэрийн зургийн индекс"
                table.cell(i, 6).merge(table.cell(i, 7))
                row_cells[8].text = str(row[2]) if row[2] is not None else ''
                table.cell(i, 8).merge(table.cell(i, 9))
            elif i == 8:
                row_cells[4].text = '1:25000 зурагт'
                row_cells[7].text = "1:100 000 зурагт"
                table.cell(i, 4).merge(table.cell(i, 6))
                table.cell(i, 7).merge(table.cell(i, 9))
            elif i == 9:
                row_cells[7].text = str(row[10]) if row[10] is not None else ''
                table.cell(i, 4).merge(table.cell(i, 6))
                table.cell(i, 7).merge(table.cell(i, 9))
            elif i == 10:
                paragraph = row_cells[4].paragraphs[0]
                value = str(row[20]) if row[20] is not None else ''
                run = paragraph.add_run(value)
                run.font.color.rgb = RGBColor(255, 0, 0)
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 13:
                paragraph = row_cells[4].paragraphs[0]
                run = paragraph.add_run()
                if export_path:
                    print("export_path found", export_path)
                    try:
                        run.add_picture(export_path, width=Inches(3.0), height=Inches(3.0))
                    except UnexpectedEndOfFileError:
                        print(f"Error: The export image file at {export_path} is corrupted or malformed.")
                    except Exception as e:
                        print(f"An error occurred: {e}")
                table.cell(i, 4).merge(table.cell(i, 9))
            elif i == 12:
                # row_cells[0].text = str(static_value) if static_value is not None else ''
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run(str(static_value.get("first", "")))
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(i, 0).merge(table.cell(i, 9))
            else:
                row_cells[4].text = str(static_value.get("second", ""))
                table.cell(i, 4).merge(table.cell(i, 9))

        doc.add_page_break()

    # Save the Word document to a temporary location
    doc_path = 'Huviin_hereg.docx'
    doc.save(doc_path)
    logger.info("=== Хувийн хэрэг амжилттай хадгалагдлаа: %s ===", doc_path)
    return doc_path

def handle_name_request_file(f):
    wb = openpyxl.load_workbook(f)
    ws = wb.active
    doc = Document()
    static_values = [
        {
            "id": "1.",
            "first": "Хүсэлт /өргөдөл/ гаргагчийн мэдээлэл /Иргэн, Аж ахуйн нэгж, Төрийн байгууллага, Төрийн бус байгууллага болон бусад/",
            "second": "Овог, нэр: Цэрэндорж Орлого\nРД: ОА92103007\nОршин суугаа хаяг: \nУтас: 99913033\nФакс: \nИ-мэйл: orlogo.ts@gazar.gov.mn\nГарын үсэг:\nХавсралт баримтын хуудасны тоо: ____________\nОгноо: _______________"
        },
        {
            "id": "2.",
            "first": "Санал болгож буй газар зүйн нэр",
            "second": "",
            "third": "1 дэх нэр",
            "fourth": "2 дахь нэр"
        },
        {
            "id": "2.",
            "first": "Санал болгож буй газар зүйн нэр",
            "second": "",
            "third": "1 дэх нэр",
            "fourth": "2 дахь нэр"
        },
        {
            "id": "3.",
            "first": "Нэрний гарал үүсэл",
            "second": "Ο Шинээр бий болсон газар зүйн объект\nΟ Газар зүйн уламжлалт нэр",
        },
        {
            "id": "4.",
            "first": "Дэвсгэр нэр /ам, булаг, гол, нуур, уул... гэх мэт/",
            "second": "",
        },
        {
            "id": "5.",
            "first": "Аймаг, нийслэл, сум, дүүрэг, баг, хорооны нэр, дугаар",
            "second": "",
        },
        {
            "id": "6.",
            "first": "Хамгийн ойр орших хот, суурин газраас алслагдах зай, километрээр /аль зүгт байрлахыг тодорхой бичих/.",
            "second": "",
        },
        {
            "id": "7.",
            "first": "Газар зүйн нэрийн солбицол /градус, минут, секунд/",
            "second": "",
        },
        {
            "id": "8.",
            "first": "Шинээр бий болсон объектод өгөх нэр, уламжлалт газар зүйн нэрийн хэрэглэгдэж буй хугацаа /жилээр/",
            "second": "Ο 50-иас дээш жил /хуучин нэр/\nΟ 10-50 хүртэлх жил /харьцангуй хуучин нэр/\nΟ 10 хүртэлх жил /шинэ нэр/",
        },
        {
            "id": "9.",
            "first": "Нэрийн талаар мэдээллээр хангагч иргэн, хуулийн этгээдийн мэдээлэл",
            "second": "Овог, нэр: Б.Сэргэлэн\nРегистрийн дугаар: БД60051471\nХаяг: Дундговь аймаг, Луус сум 1-р баг Наран\nУтас:  88261547\nИ-мэйл: sergelenb@gmail.com",
        },
        {
            "id": "10.",
            "first": "Эрх бүхий байгууллага болон орон нутгийн зөвлөлийн зөвлөмж",
            "second": "1.Сумын ГЗНСЗ-ийн хурлын шийдвэр\n2.Сумын ИТХ-ын тогтоол\n3.Аймгийн ГЗНСЗ-ийн хурлын шийдвэр\n4.Аймгийн ИТХ-ын тогтоол\n5.ГЗБГЗЗГ-ын ГЗНЗ-ийн хурлын шийдвэр\n6.Газар зүйн нэрийн Үндэсний зөвлөлийн зөвлөмж\n7.Засгийн газар\n8.Үндэсний аюулгүй байдлын зөвлөл\n9.Улсын Их Хурлын тогтоол",
        },
        {
            "id": "11.",
            "first": "Гэрэл зураг",
            "second": "",
            "third": "",
            "description": "Тайлбар: 1-8 ширхэг гэрэл зураг оруулах /зураг дарсан зүг, чиг бичих/;\n/Жишээ нь: Зүүн зүгээс эсвэл Зүүн урд зүгээс гэх мэтээр бичнэ/.",
        },
        {
            "id": "11.",
            "first": "Гэрэл зураг",
            "second": "",
            "third": "",
            "description": "Тайлбар: 1-8 ширхэг гэрэл зураг \nоруулах /зураг дарсан зүг, чиг\nбичих/;\n/Жишээ нь: Зүүн зүгээс эсвэл \nЗүүн урд зүгээс гэх мэтээр \nбичнэ/.",
        },
        {
            "id": "12.",
            "first": "Байршлын зураг",
            "second": "",
            "description": "/Газар зүйн нэрийн зураг болон сумын бүдүүвч зураг дээр харагдах байдал/"
        },
        {
            "id": "12.",
            "first": "Байршлын зураг",
            "second": "",
            "description": "\n\n\n/Газар зүйн нэрийн зураг болон сумын бүдүүвч зураг дээр харагдах байдал/"
        },
    ]

    for row in ws.iter_rows(min_row=3, values_only=True):
        # doc.add_heading('', 0)
        top_right_paragraph = doc.add_paragraph('Хавсралт 1')
        top_right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = top_right_paragraph.runs[0]

        title = doc.add_paragraph('')
        run = title.add_run('Газар зүйн нэрийг шинээр өгөх хүсэлтийн маягт /өргөдөл/')
        run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        top_right_paragraph = doc.add_paragraph('Зөвхөн албан хэрэгцээнд:')
        top_right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = top_right_paragraph.runs[0]

        table = doc.add_table(rows=15, cols=15)
        table.style = 'Table Grid'
        table.autofit = False
        total_width_cm = 20  # Example total width for table
        table.columns[0].width = Cm(total_width_cm * 0.4)
        for i, static_value in enumerate(static_values):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(static_value.get("id", ""))
            row_cells[1].text = str(static_value.get("first", ""))
            if i == 1:
                horizontal = table.cell(i, 1).merge(table.cell(i, 3))
                horizontal.merge(table.cell(i+1, 3))
                table.cell(i, 0).merge(table.cell(i+1, 0))
                row_cells[4].text = str(static_value.get("third", ""))
                table.cell(i, 4).merge(table.cell(i, 5))
                table.cell(i + 1 , 4).merge(table.cell(i+1, 5))
            elif i == 2:
                row_cells[4].text = str(static_value.get("fourth", ""))
            elif i == 11 or i == 13:
                table.cell(i, 0).merge(table.cell(i+1, 0))
                table.cell(i, 1).merge(table.cell(i, 14))
                row_cells[1].text = str(static_value.get("first", ""))
            elif i == 12 or i == 14:
                table.cell(i, 1).merge(table.cell(i, 14))
                row_cells[1].text = str(static_value.get("description", ""))
            else:
                table.cell(i, 1).merge(table.cell(i, 5)) 
            
            if i == 7:
                longitude1 = "1.Өргөрөг: " + ("" if not row[14] else str(row[14]))
                latitude1 = "\nУртраг: " + ("" if not row[15] else str(row[15]))

                longitude2 = "\n2.Өргөрөг: " + ("" if not row[16] else str(row[16]))
                latitude2 = "\nУртраг: " + ("" if not row[17] else str(row[17]))
                coordinate = longitude1 + latitude1
                if row[16]:
                    coordinate = coordinate + longitude2 + latitude2
                row_cells[6].text = coordinate
            elif i == 1:
                row_cells[6].text = (str(row[3]) if str(row[3]) is not None else "")
            elif i == 4:
                row_cells[6].text = (str(row[5]) if str(row[5]) is not None else "")
            elif i == 5:
                row_cells[6].text = str(row[18]) if row[18] is not None else ''
            elif i == 6:
                row_cells[6].text = str(row[19]) if row[19] is not None else ''
            else:
                row_cells[6].text = str(static_value.get("second", ""))
            table.cell(i, 6).merge(table.cell(i, 14))

        doc.add_page_break()

    # Save the Word document to a temporary location
    doc_path = 'huseltiing_maygt.docx'
    doc.save(doc_path)
    return doc_path

def home(request):
    return render(request, 'base.html')

@csrf_exempt
def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                doc_path = handle_uploaded_file(request.FILES['file'])
                with open(doc_path, 'rb') as fh:
                    response = HttpResponse(fh.read(),
                                            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    response['Content-Disposition'] = 'inline; filename=' + doc_path
                    return response
            except ValueError as e:
                logger.error("Validation алдаа: %s", str(e))
                return JsonResponse({'error': str(e)}, status=400)
            except Exception as e:
                logger.exception("Тодорхойгүй алдаа:")
                return JsonResponse({'error': 'Серверийн алдаа: ' + str(e)}, status=500)
    else:
        form = UploadFileForm()
    return render(request, 'fileconverter/upload.html', {'form': form})

@csrf_exempt
def name_request(request):
    if request.method == 'POST':
        form = ChangeRequestFileForm(request.POST, request.FILES)
        if form.is_valid():
            doc_path = handle_name_request_file(request.FILES['file'])
            with open(doc_path, 'rb') as fh:
                response = HttpResponse(fh.read(),
                                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                response['Content-Disposition'] = 'inline; filename=' + doc_path
                return response
    else:
        form = ChangeRequestFileForm()
    return render(request, 'fileconverter/name_request.html', {'form': form})

